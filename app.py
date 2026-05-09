## Create the Streamlit app

import os
import json
import numpy as np
import faiss
import fitz
import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from pydantic import BaseModel
from typing import Optional
import openpyxl

class CrashPrediction(BaseModel):
    state_case_number: str
    manner_of_crash: str
    first_harmful_event: str
    v1_maneuver: str
    v2_maneuver: Optional[str] = None
    v3_maneuver: Optional[str] = None
    v4_maneuver: Optional[str] = None
    reasoning: str

# Load environment
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# --- Load FAISS index and metadata ---
@st.cache_resource
def load_rag():
    index = faiss.read_index(r"C:\Users\Hamidreza\OneDrive\Desktop\LLM\LLM_Project\ecrash_faiss.index")
    with open(r"C:\Users\Hamidreza\OneDrive\Desktop\LLM\LLM_Project\ecrash_metadata.json", "r") as f:
        metadata = json.load(f)
    return index, metadata

index, metadata = load_rag()

# --- RAG retrieval function ---
def retrieve_relevant_chunks(query, top_k=5):
    response = client.embeddings.create(
        model="text-embedding-3-small",
        input=[query]
    )
    query_embedding = np.array([response.data[0].embedding], dtype="float32")
    distances, indices = index.search(query_embedding, top_k)
    return [metadata[idx] for idx in indices[0]]

# --- System prompt ---
SYSTEM_PROMPT = """You are a Louisiana crash report coding specialist trained on the eCrash User Guide.
Given a crash narrative, you must classify the crash by assigning official codes to these fields:

1. Manner of Crash (MOC): How the vehicles came together
2. First Harmful Event (FHE): The first event that caused injury or damage
3. Vehicle Maneuver: What each vehicle was doing before the crash

CRITICAL RULES FOR PARKED vs STOPPED:
- 500 Parked = vehicle is in a designated parking area, parking lot, or off the roadway. NOT in transport.
- 501 Stopped = vehicle is stopped in a TRAVEL LANE. IS in transport.
- THE LOCATION DETERMINES THE CODE: If any part of the vehicle or its attached trailer is in a travel lane, the vehicle is STOPPED (501), not PARKED (500).
- A vehicle that is STOPPED (501) in a travel lane IS considered "in transport" for MOC purposes.

CRITICAL RULES FOR MANNER OF CRASH (MOC):
- FIRST determine how many vehicles are "in transport." If ONLY ONE vehicle was in transport, MOC MUST be 000. Do NOT consider point of contact if only one vehicle was in transport.
- ONLY if TWO OR MORE vehicles were in transport, then determine MOC from the ACTUAL POINT OF CONTACT:
  - 300 Front to rear = FRONT of one vehicle strikes the REAR of another, both facing same direction
  - 505 Sideswipe same direction = SIDE of one vehicle contacts the SIDE of another, including hitting an open door or the side of a trailer
  - 105 Angle perpendicular = vehicles are at an angle to each other, including a vehicle perpendicular to traffic being struck

RULES FOR FIRST HARMFUL EVENT (FHE):
- If a moving vehicle strikes another vehicle "in transport" (including stopped in travel lane), FHE = 201
- If a moving vehicle strikes a PARKED vehicle (off roadway), FHE = 202

RULES FOR VEHICLE MANEUVER:
- Determine what each vehicle was doing IMMEDIATELY BEFORE the crash
- Do NOT code a vehicle as turning unless the narrative clearly states it was turning at the moment of impact

IMPORTANT:
- Use ONLY the official code values listed below
- Output the code number AND description
- Output valid JSON only, no extra text

CODE VALUES:

Manner of Crash (MOC):
000~Not a collision between two motor vehicles in transport
100~Angle - left overtake
101~Angle - left opposite direction
102~Angle - left into flow
103~Angle - right into flow
104~Angle - right overtake
105~Angle - perpendicular/other angle
200~Front to front - head on
300~Front to rear - rear end
400~Backing - rear to front
401~Backing - rear to rear
402~Backing - rear to side
500~Angle - left across flow
501~Angle - right across flow
502~Sideswipe - opposite direction
505~Sideswipe - same direction
980~Other
999~Unknown

First Harmful Event (FHE):
100~Cargo/equipment loss or shift
101~Fell/jumped from motor vehicle
102~Fire/explosion
103~Immersion, full or partial
104~Jackknife
105~Overturn/rollover
106~Thrown or falling object
198~Other non-collision harmful event
200~Collision with animal (live)
201~Collision with motor vehicle in transport
202~Collision with parked motor vehicle
203~Collision with pedalcycle
204~Collision with pedestrian
205~Collision with railway vehicle
206~Collision with object at rest from MV in transport
207~Collision with falling, shifting cargo
208~Collision with work zone/maintenance equipment
209~Collision with farm equipment
297~Collision with other non-motorist
298~Collision with other non-fixed object
300~Collision with bridge overhead structure
301~Collision with bridge pier or support
302~Collision with bridge rail
303~Collision with cable barrier
304~Collision with concrete traffic barrier
305~Collision with culvert
306~Collision with curb
307~Collision with ditch
308~Collision with embankment
309~Collision with fence
310~Collision with guardrail end terminal
311~Collision with guardrail face
312~Collision with impact attenuator/crash cushion
313~Collision with mailbox
314~Collision with traffic sign support
315~Collision with traffic signal support
316~Collision with tree (standing)
317~Collision with utility pole/light support
396~Collision with other post, pole, or support
397~Collision with other traffic barrier
398~Collision with other fixed object
399~Collision with unknown fixed object

Vehicle Maneuver:
100~Going straight
101~Backing
102~Merging
103~Making U-turn
104~Negotiating a curve
106~Turning left
107~Turning right
108~Traveling wrong way
200~Leaving a parking position
400~Slowing
500~Parked
501~Stopped
980~Other
999~Unknown"""

# --- Interpret function ---
def interpret_narrative(narrative, case_number):
    rag_results = retrieve_relevant_chunks("manner of crash first harmful event vehicle maneuver coding rules", top_k=5)
    coding_rules = "\n\n---\n\n".join([r["text"] for r in rag_results])
    
    user_prompt = f"""## Relevant Coding Rules from eCrash User Guide:
{coding_rules}

## Crash Narrative (State Case #{case_number}):
{narrative}

Think step by step:
1. How many vehicles are involved?
2. For EACH vehicle: Is any part of it (including attached trailers) in a travel lane? If yes = in transport. If in a parking lot or off the roadway = parked, not in transport.
3. How many vehicles are "in transport"? If only ONE, MOC must be 000.
4. What was each vehicle doing IMMEDIATELY BEFORE the crash?
5. What was the FIRST harmful event?
6. ONLY if two or more vehicles are in transport: What is the ACTUAL POINT OF CONTACT?

Then output ONLY valid JSON in this format:
{{
  "state_case_number": "{case_number}",
  "manner_of_crash": "CODE~DESCRIPTION",
  "first_harmful_event": "CODE~DESCRIPTION",
  "v1_maneuver": "CODE~DESCRIPTION",
  "v2_maneuver": "CODE~DESCRIPTION or DELETE or null",
  "v3_maneuver": "CODE~DESCRIPTION or DELETE or null",
  "v4_maneuver": "CODE~DESCRIPTION or DELETE or null",
  "reasoning": "Your step-by-step reasoning here"
}}"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.2,
        response_format={"type": "json_object"}
    )
    
    raw_result = json.loads(response.choices[0].message.content)
    
    try:
        validated = CrashPrediction(**raw_result)
        return validated.model_dump()
    except Exception as e:
        st.warning(f"Validation error: {e}")
        return raw_result

# --- Streamlit UI ---
st.set_page_config(page_title="Crash Narrative Interpreter", layout="wide")
st.title("RAG-Augmented Crash Narrative Interpreter")
st.caption("Powered by GPT-4o + FAISS + eCrash User Guide")

tab1, tab2 = st.tabs(["Single Narrative", "Batch Processing"])

# --- Tab 1: Single narrative input ---
with tab1:
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Input")
        case_number = st.text_input("State Case Number", placeholder="e.g., 2025020863")
        
        narratives_dir = r"C:\Users\Hamidreza\OneDrive\Desktop\LLM\LLM_Project\narratives"
        
        narrative_text = ""
        if case_number:
            filepath = os.path.join(narratives_dir, f"{case_number}.docx")
            if os.path.exists(filepath):
                doc = Document(filepath)
                narrative_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                st.success(f"Found narrative file: {case_number}.docx")
                st.text_area("Narrative Preview", narrative_text, height=300, disabled=True)
            else:
                st.error(f"No file found: {case_number}.docx")
        
        if st.button("Interpret Narrative", type="primary", disabled=not (case_number and narrative_text)):
            with st.spinner("Processing..."):
                result = interpret_narrative(narrative_text, case_number)
                st.session_state["result"] = result
    
    with col2:
        st.subheader("Output")
        if "result" in st.session_state:
            r = st.session_state["result"]
            st.metric("Manner of Crash", r["manner_of_crash"])
            st.metric("First Harmful Event", r["first_harmful_event"])
            st.metric("V1 Maneuver", r["v1_maneuver"])
            st.metric("V2 Maneuver", r.get("v2_maneuver", "N/A"))
            st.metric("V3 Maneuver", r.get("v3_maneuver", "N/A"))
            st.metric("V4 Maneuver", r.get("v4_maneuver", "N/A"))
            
            with st.expander("Reasoning"):
                st.write(r.get("reasoning", ""))
            
            with st.expander("Full JSON"):
                st.json(r)
            
            st.divider()
            
            # Confirm button - saves to Excel
            if st.button("✅ Confirm and Save to Excel", type="primary"):
                excel_path = r"C:\Users\Hamidreza\OneDrive\Desktop\LLM\LLM_Project\State_number.xlsx"
                
                try:
                    wb = openpyxl.load_workbook(excel_path)
                    ws = wb.active
                    
                    # Headers are on row 2, not row 1
                    headers = {cell.value: cell.column for cell in ws[2] if cell.value is not None}
                    
                    target_row = None
                    case_col = headers["State Case Number"]
                    
                    # Data starts from row 3
                    for row_num in range(3, ws.max_row + 1):
                        cell_val = ws.cell(row=row_num, column=case_col).value
                        if str(cell_val).strip() == str(r["state_case_number"]).strip():
                            target_row = row_num
                            break
                    
                    if target_row is None:
                        st.error(f"Case number {r['state_case_number']} not found in Excel.")
                    else:
                        ws.cell(row=target_row, column=headers["Correct Manner Code"], value=r["manner_of_crash"])
                        ws.cell(row=target_row, column=headers["First Harmful Event(FHE)"], value=r["first_harmful_event"])
                        ws.cell(row=target_row, column=headers["V1 Correct Maneuver"], value=r["v1_maneuver"])
                        ws.cell(row=target_row, column=headers["V2 Correct Maneuver"], value=r.get("v2_maneuver"))
                        ws.cell(row=target_row, column=headers["V3 Correct Maneuver"], value=r.get("v3_maneuver"))
                        ws.cell(row=target_row, column=headers["V4 Correct Maneuver"], value=r.get("v4_maneuver"))
                        
                        wb.save(excel_path)
                        st.success(f"✅ Saved to row {target_row} in Excel!")
                except PermissionError:
                    st.error("Cannot save - the Excel file is currently open. Please close it and try again.")
                except Exception as e:
                    st.error(f"Error: {e}")

# --- Tab 2: Batch processing ---
with tab2:
    st.subheader("Batch Process from Excel")
    
    excel_path = r"C:\Users\Hamidreza\OneDrive\Desktop\LLM\LLM_Project\State_number.xlsx"
    narratives_dir = r"C:\Users\Hamidreza\OneDrive\Desktop\LLM\LLM_Project\narratives"
    
    # Show case numbers from Excel
    try:
        wb = openpyxl.load_workbook(excel_path, read_only=True)
        ws = wb.active
        
        # Headers are on row 2
        headers_preview = {cell.value: cell.column for cell in ws[2] if cell.value is not None}
        case_col = headers_preview["State Case Number"]
        
        case_numbers = []
        for row_num in range(3, ws.max_row + 1):
            val = ws.cell(row=row_num, column=case_col).value
            if val:
                case_numbers.append(str(val).strip())
        wb.close()
        
        st.write(f"Found {len(case_numbers)} case numbers in Excel: {', '.join(case_numbers[:5])}{'...' if len(case_numbers) > 5 else ''}")
    except Exception as e:
        st.error(f"Cannot read Excel file: {e}")
        case_numbers = []
    
    if st.button("Process All and Save to Excel", type="primary", disabled=not case_numbers):
        try:
            wb = openpyxl.load_workbook(excel_path)
            ws = wb.active
            headers = {cell.value: cell.column for cell in ws[2] if cell.value is not None}
            
            batch_results = []
            progress = st.progress(0)
            status = st.empty()
            
            for i, case_num in enumerate(case_numbers):
                narrative_file = os.path.join(narratives_dir, f"{case_num}.docx")
                
                if not os.path.exists(narrative_file):
                    status.warning(f"⚠️ No narrative file for {case_num}, skipping")
                    progress.progress((i + 1) / len(case_numbers))
                    continue
                
                status.info(f"Processing {case_num}...")
                
                doc_file = Document(narrative_file)
                narrative = "\n".join([p.text for p in doc_file.paragraphs if p.text.strip()])
                
                result = interpret_narrative(narrative, case_num)
                batch_results.append(result)
                
                # Find matching row in Excel
                for row_num in range(3, ws.max_row + 1):
                    cell_val = ws.cell(row=row_num, column=headers["State Case Number"]).value
                    if str(cell_val).strip() == case_num:
                        ws.cell(row=row_num, column=headers["Correct Manner Code"], value=result["manner_of_crash"])
                        ws.cell(row=row_num, column=headers["First Harmful Event(FHE)"], value=result["first_harmful_event"])
                        ws.cell(row=row_num, column=headers["V1 Correct Maneuver"], value=result["v1_maneuver"])
                        ws.cell(row=row_num, column=headers["V2 Correct Maneuver"], value=result.get("v2_maneuver"))
                        ws.cell(row=row_num, column=headers["V3 Correct Maneuver"], value=result.get("v3_maneuver"))
                        ws.cell(row=row_num, column=headers["V4 Correct Maneuver"], value=result.get("v4_maneuver"))
                        break
                
                progress.progress((i + 1) / len(case_numbers))
            
            wb.save(excel_path)
            st.session_state["batch_results"] = batch_results
            status.success(f"✅ Processed {len(batch_results)} narratives and saved to Excel!")
        except PermissionError:
            st.error("Cannot save - the Excel file is currently open. Please close it and try again.")
        except Exception as e:
            st.error(f"Error: {e}")
    
    if "batch_results" in st.session_state:
        st.divider()
        st.subheader("Results")
        for r in st.session_state["batch_results"]:
            with st.expander(f"Case #{r['state_case_number']}"):
                st.metric("Manner of Crash", r["manner_of_crash"])
                st.metric("First Harmful Event", r["first_harmful_event"])
                st.metric("V1 Maneuver", r["v1_maneuver"])
                st.metric("V2 Maneuver", r.get("v2_maneuver", "N/A"))
                st.metric("V3 Maneuver", r.get("v3_maneuver", "N/A"))
                st.json(r)